from flask import Flask, request, jsonify, render_template
import os
import tempfile
import pypdf2
from docx import Document
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document as LC_Document
from langchain_astradb import AstraDBVectorStore

app = Flask(__name__)

# Astra DB Configuration (Replace with actual values)
ASTRA_DB_APPLICATION_TOKEN = "AstraCS:QNlafqYWyloiLqCESgecoTFe:dfc043c7cbb8c99cfc302f074a1f721f1f2d658e8e7350b9abc6235f81005360"
ASTRA_DB_API_ENDPOINT = "https://ebc8f41b-62ae-449f-b6bb-36284a0bee58-us-east-2.apps.astra.datastax.com"
ASTRA_DB_COLLECTION = "filesdata"

# Initialize AstraDB Vector Store
vector_store = AstraDBVectorStore(
    token=ASTRA_DB_APPLICATION_TOKEN,
    api_endpoint=ASTRA_DB_API_ENDPOINT,
    collection_name=ASTRA_DB_COLLECTION,
    autodetect_collection=True
)

def extract_text(file_path, file_type):
    """Extract text from different file formats"""
    text = ""
    
    if file_type == "pdf":
        with pypdf2.open(file_path) as pdf:
            text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    
    elif file_type == "docx":
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    
    else:  # Default: plain text
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    
    return text

@app.route('/')
def home():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    # Determine file extension
    ext = file.filename.split(".")[-1].lower()
    if ext not in ["txt", "pdf", "docx"]:
        return jsonify({"error": "Unsupported file type"}), 400
    
    # Save temporary file
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, file.filename)
    file.save(file_path)
    
    # Extract text based on file type
    text = extract_text(file_path, ext)
    os.remove(file_path)

    # Split text into chunks
    text_splitter = CharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separator="\n"
    )
    chunks = text_splitter.split_text(text)
    
    # Store in Astra DB
    documents = [LC_Document(page_content=chunk, metadata={"source": file.filename}) for chunk in chunks]
    vector_store.add_documents(documents)
    
    return jsonify({"message": "File processed and stored successfully", "chunks": len(chunks)})

from flask import Flask

def handler(event, context):
    return app(event, context)

